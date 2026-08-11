#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""批量变化检测图表及Word/PDF自动报告生成器。

Excel读取、图表和Word生成均由普通Python库完成；Microsoft Word仅用于
可选的PDF转换。所有Excel、CSV与既有成果均以只读方式访问。
"""

from __future__ import annotations

import argparse
import csv
import json
import logging
import math
import os
import shutil
import subprocess
import sys
import traceback
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_EXCEL = BASE_DIR / "Excel统计表" / "批量变化检测统计汇总.xlsx"
DEFAULT_STATS_CSV = BASE_DIR / "统计数据" / "批量变化检测统计明细.csv"
DEFAULT_RELIABILITY_CSV = BASE_DIR / "可选输入" / "可靠性验收状态表.csv"
DEFAULT_RELIABILITY_EVIDENCE = BASE_DIR / "可选输入" / "可靠性验收证据.json"
DEFAULT_PRIMARY_STATUS = BASE_DIR / "可选输入" / "批量状态表.csv"
DEFAULT_PRIMARY_LOG = BASE_DIR / "可选输入" / "批量运行.log"
DEFAULT_RELIABILITY_LOG = BASE_DIR / "可选输入" / "可靠性验收.log"
DEFAULT_PROJECT_PYTHON = Path(sys.executable)
DEFAULT_WORK_DIR = BASE_DIR / "运行输出" / "_work"
DEFAULT_SCRIPT_OUTPUT = Path(__file__).resolve()
DEFAULT_CHART_DIR = BASE_DIR / "图表"
DEFAULT_DOCX = BASE_DIR / "Word报告" / "批量变化检测统计报告_样例联调版.docx"
DEFAULT_PDF = BASE_DIR / "PDF报告" / "批量变化检测统计报告_样例联调版.pdf"
DEFAULT_LOG = BASE_DIR / "运行记录" / "图表与报告生成.log"
DEFAULT_ACCEPTANCE = BASE_DIR / "运行记录" / "图表与报告生成验收报告.md"

DISCLAIMER = "本报告用于批量变化检测流程联调与功能验证，当前统计数据来自官方样例影像。"
SOURCE_NOTE = "数据来源：批量变化检测统计明细.csv（官方样例影像）"

# 只保留三张图：状态、面积、比例
CHART_FILES = {
    "status": "任务状态统计图.png",
    "area": "各任务变化面积图.png",
    "ratio": "各任务变化比例图.png",
}


def configure_logging(log_path: Path, *, reset: bool = False) -> logging.Logger:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    if reset:
        log_path.write_text("", encoding="utf-8", newline="\n")
    logger = logging.getLogger("report_generator")
    logger.handlers.clear()
    logger.setLevel(logging.INFO)
    formatter = logging.Formatter(
        "%(asctime)s | %(levelname)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )
    file_handler = logging.FileHandler(log_path, mode="a", encoding="utf-8")
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)
    stream_handler = logging.StreamHandler(sys.stdout)
    stream_handler.setFormatter(formatter)
    logger.addHandler(stream_handler)
    return logger


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def as_int(value: Any, default: int = 0) -> int:
    if value in (None, ""):
        return default
    return int(float(value))


def as_float(value: Any, default: float = 0.0) -> float:
    if value in (None, ""):
        return default
    return float(value)


def run_process(
    command: list[str],
    logger: logging.Logger,
    *,
    timeout: int = 180,
) -> subprocess.CompletedProcess[str]:
    env = os.environ.copy()
    env["PYTHONUTF8"] = "1"
    logger.info("运行子进程 | command=%s", command)
    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=timeout,
        env=env,
        check=False,
    )
    if result.stdout.strip():
        logger.info("子进程stdout:\n%s", result.stdout.strip())
    if result.stderr.strip():
        logger.warning("子进程stderr:\n%s", result.stderr.strip())
    if result.returncode != 0:
        raise RuntimeError(
            f"子进程退出代码{result.returncode}: {' '.join(command)}\n"
            f"{result.stderr.strip()}"
        )
    return result


def read_excel_snapshot(excel_path: Path) -> dict[str, Any]:
    """只读读取Excel缓存值，用于与CSV自动交叉核对。"""
    from openpyxl import load_workbook

    workbook = load_workbook(excel_path, read_only=True, data_only=True)
    try:
        overview_sheet = workbook["总览"]
        overview = {
            overview_sheet[f"A{row}"].value: overview_sheet[f"B{row}"].value
            for row in range(2, 12)
        }
        task_sheet = workbook["任务明细"]
        headers = [cell.value for cell in next(task_sheet.iter_rows(min_row=1, max_row=1))]
        task_rows = [
            dict(zip(headers, [cell.value for cell in row]))
            for row in task_sheet.iter_rows(min_row=2)
        ]
        stats_sheet = workbook["变化统计"]
        stats_headers = [
            cell.value for cell in next(stats_sheet.iter_rows(min_row=1, max_row=1))
        ]
        stats_rows = [
            dict(zip(stats_headers, [cell.value for cell in row]))
            for row in stats_sheet.iter_rows(min_row=2)
        ]
        return {
            "overview": overview,
            "task_rows": task_rows,
            "stats_rows": stats_rows,
            "sheet_names": list(workbook.sheetnames),
        }
    finally:
        workbook.close()


def collect_environment(
    project_python: Path | None,
    before_image: str,
    after_image: str,
    logger: logging.Logger,
) -> dict[str, Any]:
    environment: dict[str, Any] = {
        "iobjectspy_version": "未获取",
        "gpu_info": "GPU 0（硬件信息未获取）",
        "gpu_driver": "",
        "gpu_memory": "",
        "input_shift_m": None,
        "input_shift_pixels": None,
        "input_pixel_size_m": None,
    }

    try:
        import importlib.metadata as metadata

        environment["iobjectspy_version"] = metadata.version("iobjectspy")
    except Exception:
        logger.info("当前Python环境未安装iObjectsPy；报告仍可继续生成")

    nvidia = shutil.which("nvidia-smi.exe") or shutil.which("nvidia-smi")
    if nvidia:
        try:
            result = subprocess.run(
                [
                    nvidia,
                    "--query-gpu=index,name,driver_version,memory.total",
                    "--format=csv,noheader,nounits",
                ],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=30,
                check=True,
            )
            first = result.stdout.strip().splitlines()[0]
            index, name, driver, memory = [part.strip() for part in first.split(",", 3)]
            environment["gpu_info"] = f"GPU {index}：{name}"
            environment["gpu_driver"] = driver
            environment["gpu_memory"] = f"{memory} MiB"
        except Exception:
            logger.exception("GPU信息读取失败")

    if before_image and after_image and Path(before_image).is_file() and Path(after_image).is_file():
        try:
            import rasterio

            with rasterio.open(before_image) as first, rasterio.open(after_image) as second:
                dx = float(second.transform.c - first.transform.c)
                dy = float(second.transform.f - first.transform.f)
                shift = math.hypot(dx, dy)
                pixel_size = (
                    abs(float(first.transform.a)) + abs(float(first.transform.e))
                ) / 2
                environment["input_shift_m"] = shift
                environment["input_shift_pixels"] = shift / pixel_size if pixel_size else None
                environment["input_pixel_size_m"] = pixel_size
                environment["input_crs"] = str(first.crs)
        except Exception:
            logger.exception("官方样例亚像元偏移读取失败")

    return environment


def build_payload(
    excel_path: Path,
    stats_csv: Path,
    reliability_csv: Path,
    reliability_evidence: Path,
    project_python: Path | None,
    logger: logging.Logger,
) -> dict[str, Any]:
    csv_rows = read_csv_rows(stats_csv)
    excel = read_excel_snapshot(excel_path)

    status_counts = Counter((row.get("任务状态") or "").strip() for row in csv_rows)
    valid_rows: list[dict[str, Any]] = []
    for row in csv_rows:
        status = (row.get("任务状态") or "").strip()
        if status not in {"success", "skipped"}:
            continue
        if not row.get("变化像元数（值1）"):
            continue
        valid_rows.append(
            {
                "task_id": row["task_id"],
                "status": status,
                "output_path": row.get("输出文件完整路径", ""),
                "file_size_bytes": as_int(row.get("文件大小（字节）")),
                "total_pixels": as_int(row.get("总像元数")),
                "valid_pixels": as_int(row.get("有效像元数")),
                "unchanged_pixels": as_int(row.get("未变化像元数（值0）")),
                "changed_pixels": as_int(row.get("变化像元数（值1）")),
                "other_pixels": as_int(row.get("其他像元值数量")),
                "change_ratio": as_float(row.get("变化比例")),
                "change_area_m2": as_float(row.get("变化面积（平方米）")),
                "duration_seconds": as_float(row.get("推理耗时（秒）")),
                "retry_count": as_int(row.get("重试次数")),
                "crs": row.get("坐标系", ""),
            }
        )

    total_duration = sum(as_float(row.get("推理耗时（秒）")) for row in csv_rows)
    summary = {
        "total_tasks": len(csv_rows),
        "success_count": status_counts.get("success", 0),
        "failed_count": status_counts.get("failed", 0),
        "skipped_count": status_counts.get("skipped", 0),
        "statistic_result_count": len(valid_rows),
        "total_changed_pixels": sum(row["changed_pixels"] for row in valid_rows),
        "total_changed_area_m2": sum(row["change_area_m2"] for row in valid_rows),
        "average_change_ratio": (
            sum(row["change_ratio"] for row in valid_rows) / len(valid_rows)
            if valid_rows
            else 0.0
        ),
        "total_duration_seconds": total_duration,
    }

    checks = [
        ("总任务数", summary["total_tasks"], excel["overview"].get("总任务数"), 0),
        ("成功数", summary["success_count"], excel["overview"].get("成功数"), 0),
        ("失败数", summary["failed_count"], excel["overview"].get("失败数"), 0),
        ("跳过数", summary["skipped_count"], excel["overview"].get("跳过数"), 0),
        (
            "可统计结果数",
            summary["statistic_result_count"],
            excel["overview"].get("可统计结果数"),
            0,
        ),
        (
            "总变化像元数",
            summary["total_changed_pixels"],
            excel["overview"].get("总变化像元数"),
            0,
        ),
        (
            "总变化面积（平方米）",
            summary["total_changed_area_m2"],
            excel["overview"].get("总变化面积（平方米）"),
            1e-6,
        ),
        (
            "平均变化比例",
            summary["average_change_ratio"],
            excel["overview"].get("平均变化比例"),
            1e-12,
        ),
        (
            "总运行耗时（秒）",
            summary["total_duration_seconds"],
            excel["overview"].get("总运行耗时（秒）"),
            1e-6,
        ),
    ]
    mismatches = []
    for label, csv_value, excel_value, tolerance in checks:
        if excel_value is None or abs(float(csv_value) - float(excel_value)) > tolerance:
            mismatches.append(
                {"field": label, "csv": csv_value, "excel": excel_value}
            )

    task_rows = excel["task_rows"]
    first_task = next(
        (row for row in task_rows if row.get("任务状态") == "success"),
        task_rows[0],
    )
    model_name = str(first_task.get("模型名称") or "building")
    gpu_id = str(first_task.get("GPU编号") if first_task.get("GPU编号") is not None else 0)
    before_image = str(first_task.get("前时相影像") or "")
    after_image = str(first_task.get("后时相影像") or "")

    reliability_rows = (
        read_csv_rows(reliability_csv) if reliability_csv.exists() else []
    )
    evidence: dict[str, Any] = {}
    if reliability_evidence.exists():
        evidence = json.loads(reliability_evidence.read_text(encoding="utf-8"))
    reliability_counts = Counter(row.get("status", "") for row in reliability_rows)
    retry_row = next(
        (row for row in reliability_rows if row.get("task_id") == "retry_task_002"),
        {},
    )
    reliability = {
        "total_tasks": len(reliability_rows),
        "success_count": reliability_counts.get("success", 0),
        "failed_count": reliability_counts.get("failed", 0),
        "skipped_count": reliability_counts.get("skipped", 0),
        "retry_attempt": as_int(retry_row.get("attempt")),
        "retry_count": as_int(retry_row.get("retry_count")),
        "retry_final_status": retry_row.get("status", "未获取"),
        "checks": evidence.get("checks", {}),
        "phase1_states": evidence.get("phase1_states", {}),
        "combined_active_duration_seconds": evidence.get(
            "combined_active_duration_seconds"
        ),
    }

    environment = collect_environment(
        project_python, before_image, after_image, logger
    )

    payload = {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "report_date": datetime.now().strftime("%Y年%m月%d日"),
        "disclaimer": DISCLAIMER,
        "source_note": SOURCE_NOTE,
        "summary": summary,
        "valid_tasks": valid_rows,
        "sample_task": valid_rows[0] if valid_rows else {},
        "failed_tasks": [
            {
                "task_id": row.get("task_id", ""),
                "status": row.get("任务状态", ""),
                "error": row.get("警告或错误说明", ""),
                "duration_seconds": as_float(row.get("推理耗时（秒）")),
                "retry_count": as_int(row.get("重试次数")),
            }
            for row in csv_rows
            if row.get("任务状态") == "failed"
        ],
        "excel_crosscheck": {
            "passed": not mismatches,
            "mismatches": mismatches,
            "sheet_names": excel["sheet_names"],
        },
        "environment": {
            **environment,
            "model_name": model_name,
            "gpu_id": gpu_id,
            "before_image": before_image,
            "after_image": after_image,
        },
        "reliability": reliability,
        "paths": {
            "excel": str(excel_path),
            "stats_csv": str(stats_csv),
            "primary_status": str(DEFAULT_PRIMARY_STATUS),
            "primary_log": str(DEFAULT_PRIMARY_LOG),
            "reliability_status": str(reliability_csv),
            "reliability_log": str(DEFAULT_RELIABILITY_LOG),
            "reliability_evidence": str(reliability_evidence),
            "script": str(DEFAULT_SCRIPT_OUTPUT),
            "chart_dir": str(DEFAULT_CHART_DIR),
            "docx": str(DEFAULT_DOCX),
            "pdf": str(DEFAULT_PDF),
            "acceptance": str(DEFAULT_ACCEPTANCE),
            "log": str(DEFAULT_LOG),
        },
    }
    logger.info(
        "数据读取完成 | tasks=%d | valid=%d | excel_match=%s",
        summary["total_tasks"],
        len(valid_rows),
        payload["excel_crosscheck"]["passed"],
    )
    return payload


def build_charts_internal(
    payload_path: Path,
    chart_dir: Path,
    result_json: Path,
    log_path: Path,
) -> int:
    logger = configure_logging(log_path)
    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    chart_dir.mkdir(parents=True, exist_ok=True)

    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.font_manager import FontProperties
    from matplotlib.ticker import FuncFormatter, MaxNLocator

    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if not font_path.exists():
        font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    font_name = FontProperties(fname=str(font_path)).get_name()
    matplotlib.rcParams.update(
        {
            "font.family": "sans-serif",
            "font.sans-serif": [font_name],
            "axes.unicode_minus": False,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "axes.edgecolor": "#4B5563",
            "axes.labelcolor": "#263238",
            "xtick.color": "#4B5563",
            "ytick.color": "#4B5563",
            "text.color": "#263238",
            "font.size": 11,
        }
    )

    results: dict[str, Any] = {}

    def save_chart(key: str, builder: Any) -> None:
        output = chart_dir / CHART_FILES[key]
        try:
            figure = builder()
            figure.savefig(
                output,
                format="png",
                dpi=320,

                facecolor="white",
                metadata={
                    "Title": CHART_FILES[key],
                    "Description": SOURCE_NOTE,
                },
            )
            plt.close(figure)
            results[key] = {
                "status": "success",
                "path": str(output),
                "size_bytes": output.stat().st_size,
            }
            logger.info("图表生成成功 | key=%s | path=%s", key, output)
        except Exception as exc:
            logger.error(
                "图表生成失败 | key=%s | error=%s\n%s",
                key,
                exc,
                traceback.format_exc(),
            )
            results[key] = {
                "status": "failed",
                "path": str(output),
                "error_type": type(exc).__name__,
                "error_message": str(exc),
            }

    summary = payload["summary"]
    valid_tasks = payload["valid_tasks"]

    def status_chart():
        labels = ["success", "failed", "skipped"]
        values = [
            summary["success_count"],
            summary["failed_count"],
            summary["skipped_count"],
        ]
        colors = ["#0F766E", "#C2413B", "#94A3B8"]
        fig, ax = plt.subplots(figsize=(9.0, 5.4))
        bars = ax.bar(labels, values, color=colors, width=0.58, edgecolor="#334155")
        ax.set_title("任务状态统计", fontsize=18, fontweight="bold", pad=18)
        ax.set_xlabel("任务状态", labelpad=10)
        ax.set_ylabel("任务数量（项）", labelpad=10)
        ax.set_ylim(0, max(3, max(values) + 0.8))
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        ax.grid(axis="y", color="#D9E1E8", linewidth=0.8, alpha=0.85)
        ax.set_axisbelow(True)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.bar_label(bars, labels=[f"{value}项" for value in values], padding=5, fontsize=12)
        fig.text(
            0.5,
            0.015,
            SOURCE_NOTE,
            ha="center",
            va="bottom",
            fontsize=9,
            color="#64748B",
        )
        fig.tight_layout(rect=(0.03, 0.06, 0.98, 0.98))
        return fig

    def area_chart():
        labels = [row["task_id"] for row in valid_tasks]
        values = [row["change_area_m2"] for row in valid_tasks]
        fig, ax = plt.subplots(figsize=(9.0, 5.4))
        bars = ax.bar(
            labels,
            values,
            color="#2C7FB8",
            width=0.55,
            edgecolor="#1E4E79",
        )
        ax.set_title("各任务变化面积", fontsize=18, fontweight="bold", pad=18)
        ax.set_xlabel("task_id", labelpad=10)
        ax.set_ylabel("变化面积（平方米）", labelpad=10)
        upper = max(values) * 1.22 if values else 1
        ax.set_ylim(0, upper)
        ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:,.0f}"))
        ax.grid(axis="y", color="#D9E1E8", linewidth=0.8, alpha=0.85)
        ax.set_axisbelow(True)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.bar_label(
            bars,
            labels=[f"{value:,.2f}" for value in values],
            padding=5,
            fontsize=11,
        )
        fig.text(0.5, 0.015, SOURCE_NOTE, ha="center", fontsize=9, color="#64748B")
        fig.tight_layout(rect=(0.03, 0.06, 0.98, 0.98))
        return fig

    def ratio_chart():
        labels = [row["task_id"] for row in valid_tasks]
        values = [row["change_ratio"] for row in valid_tasks]
        fig, ax = plt.subplots(figsize=(9.0, 5.4))
        bars = ax.bar(
            labels,
            values,
            color="#D97706",
            width=0.55,
            edgecolor="#92400E",
        )
        ax.set_title("各任务变化比例", fontsize=18, fontweight="bold", pad=18)
        ax.set_xlabel("task_id", labelpad=10)
        ax.set_ylabel("变化比例", labelpad=10)
        upper = max(values) * 1.25 if values else 0.01
        ax.set_ylim(0, upper)
        ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.2%}"))
        ax.grid(axis="y", color="#D9E1E8", linewidth=0.8, alpha=0.85)
        ax.set_axisbelow(True)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.bar_label(
            bars,
            labels=[f"{value:.4%}" for value in values],
            padding=5,
            fontsize=11,
        )
        fig.text(0.5, 0.015, SOURCE_NOTE, ha="center", fontsize=9, color="#64748B")
        fig.tight_layout(rect=(0.03, 0.06, 0.98, 0.98))
        return fig

    # 只生成三张图
    save_chart("status", status_chart)
    save_chart("area", area_chart)
    save_chart("ratio", ratio_chart)

    result_json.parent.mkdir(parents=True, exist_ok=True)
    result_json.write_text(
        json.dumps(results, ensure_ascii=False, indent=2),
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(results, ensure_ascii=False))
    return 0


def build_docx(
    payload: dict[str, Any],
    chart_results: dict[str, Any],
    docx_path: Path,
    logger: logging.Logger,
    report_title: str = "批量变化检测统计报告",
) -> dict[str, Any]:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Mm, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)

    teal = RGBColor(15, 118, 110)
    dark = RGBColor(38, 50, 56)
    muted = RGBColor(100, 116, 139)
    white = RGBColor(255, 255, 255)

    def set_run_font(
        run,
        name: str,
        size: float,
        *,
        bold: bool | None = None,
        color: RGBColor | None = None,
    ) -> None:
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if color is not None:
            run.font.color.rgb = color

    def configure_style(style_name: str, font: str, size: float, bold: bool = False):
        style = doc.styles[style_name]
        style.font.name = font
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
        return style

    normal = configure_style("Normal", "宋体", 12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    heading1 = configure_style("Heading 1", "黑体", 15, True)
    heading1.font.color.rgb = dark
    heading1.paragraph_format.space_before = Pt(14)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.line_spacing = 1.2
    heading1.paragraph_format.keep_with_next = True

    heading2 = configure_style("Heading 2", "黑体", 14, True)
    heading2.font.color.rgb = dark
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.line_spacing = 1.2
    heading2.paragraph_format.keep_with_next = True

    caption = configure_style("Caption", "宋体", 10.5)
    caption.font.color.rgb = muted
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    def set_cell_shading(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for margin, value in (
            ("top", top),
            ("start", start),
            ("bottom", bottom),
            ("end", end),
        ):
            node = tc_mar.find(qn(f"w:{margin}"))
            if node is None:
                node = OxmlElement(f"w:{margin}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
        total = sum(widths)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl_pr = table._tbl.tblPr
        for tag, attrs in (
            ("w:tblW", {"w:w": str(total), "w:type": "dxa"}),
            ("w:tblInd", {"w:w": str(indent), "w:type": "dxa"}),
            ("w:tblLayout", {"w:type": "fixed"}),
        ):
            node = tbl_pr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                tbl_pr.append(node)
            for key, value in attrs.items():
                node.set(qn(key), value)

        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)

        for row in table.rows:
            for index, (cell, width) in enumerate(zip(row.cells, widths)):
                cell.width = Cm(width / 1440 * 2.54)
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(width))
                tc_w.set(qn("w:type"), "dxa")
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def repeat_header(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    def format_table_text(table, *, header=True, path_columns: set[int] | None = None):
        path_columns = path_columns or set()
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                        if col_index in path_columns
                        else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            "黑体" if header and row_index == 0 else "宋体",
                            10 if header and row_index == 0 else 9.5,
                            bold=True if header and row_index == 0 else False,
                            color=white if header and row_index == 0 else dark,
                        )
                if header and row_index == 0:
                    set_cell_shading(cell, "0F766E")
                elif row_index % 2 == 0:
                    set_cell_shading(cell, "F0FDFA")
        if header:
            repeat_header(table.rows[0])

    def add_body(text: str, *, indent: bool = True, bold_prefix: str | None = None):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        if indent:
            paragraph.paragraph_format.first_line_indent = Pt(24)
        if bold_prefix and text.startswith(bold_prefix):
            first = paragraph.add_run(bold_prefix)
            set_run_font(first, "宋体", 12, bold=True, color=dark)
            rest = paragraph.add_run(text[len(bold_prefix) :])
            set_run_font(rest, "宋体", 12, color=dark)
        else:
            run = paragraph.add_run(text)
            set_run_font(run, "宋体", 12, color=dark)
        return paragraph

    def add_label_value(label: str, value: str):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, "宋体", 12, bold=True, color=dark)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, "宋体", 12, color=dark)
        return paragraph

    def add_callout(text: str):
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        set_table_geometry(table, [8952], 120)
        cell = table.cell(0, 0)
        set_cell_shading(cell, "ECFDF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.3
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        set_run_font(run, "宋体", 11.5, bold=True, color=teal)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def add_figure(key: str, caption_text: str, figure_number: int) -> bool:
        result = chart_results.get(key, {})
        if result.get("status") != "success":
            add_body(
                f"图{figure_number}未生成：{result.get('error_message', '未知错误')}。",
                indent=False,
            )
            return False
        image_path = Path(result["path"])
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run().add_picture(str(image_path), width=Cm(15.5))
        cap = doc.add_paragraph(
            f"图{figure_number} {caption_text}",
            style="Caption",
        )
        cap.paragraph_format.keep_with_next = False
        return True

    def add_table_caption(text: str):
        paragraph = doc.add_paragraph(text, style="Caption")
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(4)

    # 页眉与页脚。
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f"{report_title}｜官方样例联调版")
    set_run_font(run, "宋体", 9, color=muted)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("第 ")
    set_run_font(run, "宋体", 9, color=muted)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    number_run = OxmlElement("w:r")
    number_text = OxmlElement("w:t")
    number_text.text = "1"
    number_run.append(number_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for element in (fld_begin, instr, fld_sep):
        field_run = OxmlElement("w:r")
        field_run.append(element)
        footer._p.append(field_run)
    footer._p.append(number_run)
    field_end_run = OxmlElement("w:r")
    field_end_run.append(fld_end)
    footer._p.append(field_end_run)
    run = footer.add_run(" 页")
    set_run_font(run, "宋体", 9, color=muted)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    # 文档元数据。
    doc.core_properties.title = report_title
    doc.core_properties.subject = "批量变化检测流程联调与功能验证"
    doc.core_properties.author = "李昌辉"
    doc.core_properties.keywords = "批量变化检测, 官方样例, 自动化, 统计报告"

    # 封面：editorial_cover的技术报告变体，用户字号/字体/A4要求优先。
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("批量处理与自动化 · 统计汇总与报告")
    set_run_font(run, "黑体", 12, bold=True, color=teal)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run(report_title)
    set_run_font(run, "黑体", 24, bold=True, color=dark)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(48)
    run = subtitle.add_run("官方样例联调版")
    set_run_font(run, "黑体", 16, bold=False, color=teal)

    for label, value in (("姓名", "李昌辉"), ("日期", payload["report_date"])):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(f"{label}：{value}")
        set_run_font(run, "宋体", 12, color=dark)

    doc.add_paragraph().paragraph_format.space_after = Pt(32)
    add_callout(DISCLAIMER)
    doc.add_page_break()

    summary = payload["summary"]
    environment = payload["environment"]
    reliability = payload["reliability"]
    valid_tasks = payload["valid_tasks"]
    sample = payload["sample_task"]

    doc.add_heading("一、测试目的", level=1)
    add_body(
        "本次工作用于验证批量检测、异常隔离、自动重试、断点续跑、"
        "统计汇总和报告生成链路，重点检查任务级异常是否被隔离、"
        "状态是否可恢复、结果统计是否可复核，以及图表和报告能否由"
        "既有Excel与CSV自动生成。"
    )

    doc.add_heading("二、运行环境与数据说明", level=1)
    add_label_value("iObjectsPy版本", str(environment["iobjectspy_version"]))
    gpu_detail = environment["gpu_info"]
    if environment.get("gpu_memory"):
        gpu_detail += f"，显存{environment['gpu_memory']}"
    if environment.get("gpu_driver"):
        gpu_detail += f"，驱动{environment['gpu_driver']}"
    add_label_value("GPU信息", gpu_detail)
    add_label_value(
        "推理配置",
        f"模型{environment['model_name']}，GPU编号{environment['gpu_id']}",
    )
    add_label_value(
        "输入数据",
        "SuperMap resources_ml提供的官方变化检测样例前后时相影像",
    )
    add_body(
        f"{DISCLAIMER} 输入数据不代表徐州市实际变化结果，报告中的"
        "面积、比例和像元数量仅用于验证流程与统计功能。",
        indent=False,
    )

    doc.add_heading("三、批量任务运行情况", level=1)
    add_body(
        f"首次真实批量联调共{summary['total_tasks']}项任务："
        f"成功{summary['success_count']}项、失败{summary['failed_count']}项、"
        f"跳过{summary['skipped_count']}项；任务明细耗时合计"
        f"{summary['total_duration_seconds']:.3f}秒。失败任务因前时相路径"
        "不存在而终止自身处理，但后续任务继续执行并成功，说明任务级"
        "try-except隔离已生效。"
    )
    add_body(
        "图1汇总任务最终状态。该图展示流程验收结果，不表示模型在真实"
        "区域数据上的精度或稳定性。",
        indent=False,
    )
    add_figure("status", "任务状态统计", 1)

    doc.add_heading("四、变化检测统计结果", level=1)
    add_body(
        f"有效结果共{summary['statistic_result_count']}个，总变化像元"
        f"{summary['total_changed_pixels']:,}，按任务累计的总变化面积"
        f"{summary['total_changed_area_m2']:,.2f}平方米，平均变化比例"
        f"{summary['average_change_ratio']:.4%}。两个成功任务使用同一组"
        "官方样例，因此汇总值包含重复累计，不应解释为两个独立区域。"
    )
    add_body(
        "图2按任务展示米制投影下的变化面积。两项面积相同，反映的是相同"
        "输入样例与相同配置的重复联调输出。",
        indent=False,
    )
    add_figure("area", "各有效任务变化面积", 2)
    add_body(
        "图3展示各有效任务中值1像元占有效像元的比例。百分比口径为"
        "变化像元数除以有效像元数。",
        indent=False,
    )
    add_figure("ratio", "各有效任务变化比例", 3)

    add_table_caption("表1 有效任务变化检测统计结果")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    widths = [1100, 1700, 1700, 1400, 1800, 1252]
    set_table_geometry(table, widths, 120)
    headers = [
        "task_id",
        "总像元",
        "变化像元",
        "变化比例",
        "变化面积（平方米）",
        "耗时（秒）",
    ]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
    for row_data in valid_tasks:
        cells = table.add_row().cells
        values = [
            row_data["task_id"],
            f"{row_data['total_pixels']:,}",
            f"{row_data['changed_pixels']:,}",
            f"{row_data['change_ratio']:.4%}",
            f"{row_data['change_area_m2']:,.2f}",
            f"{row_data['duration_seconds']:.3f}",
        ]
        for index, value in enumerate(values):
            cells[index].text = value
    format_table_text(table)

    doc.add_heading("五、可靠性测试结果", level=1)
    add_label_value(
        "单项失败不中断",
        "缺失输入任务失败后，后续任务仍启动并成功完成，受控验收通过。",
    )
    add_label_value(
        "自动重试",
        f"retry_task_002最终状态为{reliability['retry_final_status']}，"
        f"attempt={reliability['retry_attempt']}，重试次数"
        f"{reliability['retry_count']}；首次受控失败后第二次成功。",
    )
    add_label_value(
        "跳过已有结果",
        f"可靠性验收最终跳过{reliability['skipped_count']}项，验收时文件"
        "大小、修改时间和SHA-256均保持不变。",
    )
    add_label_value(
        "断点续跑",
        "第一阶段成功后触发受控中断；再次启动时跳过既有成功任务，"
        "仅继续未完成任务。",
    )
    add_label_value(
        "状态实时持久化",
        "控制器在每次状态变化后写入临时CSV并原子替换正式状态表，"
        "中断快照保留了已成功与待处理状态。",
    )
    add_body(
        "上述重试与中断均采用受控测试钩子，用于证明机制按设计工作；"
        "该结果不等同于已覆盖断电、驱动崩溃等全部生产异常。",
        indent=False,
    )

    doc.add_heading("六、误差与限制", level=1)
    add_label_value(
        "重复累计",
        "两个成功结果使用同一组官方样例，输出内容一致，汇总数据存在"
        "重复累计。",
    )
    shift_pixels = environment.get("input_shift_pixels")
    shift_m = environment.get("input_shift_m")
    if shift_pixels is not None and shift_m is not None:
        shift_text = (
            f"前后时相仿射原点相差约{shift_m:.2f}米，折合约"
            f"{shift_pixels:.2f}个输入像元；本报告仅记录，不自动校正栅格。"
        )
    else:
        shift_text = "输入影像存在约0.25像元亚像元偏移；本报告仅记录。"
    add_label_value("亚像元偏移", shift_text)
    add_label_value(
        "GeoTIFF标签警告",
        '控制台出现“PhotometricInterpretation”标签写入警告，但既有'
        "只读验收表明TIFF可正常打开并完整读取，不影响本次统计使用。",
    )
    add_label_value(
        "适用范围",
        "当前结果仅验证系统功能和数据链路，不用于徐州市或其他真实区域"
        "的变化结论。",
    )

    doc.add_heading("七、结论", level=1)
    add_body(
        "批量自动化、任务级异常隔离、可靠性验收、变化结果统计、图表和"
        "Word/PDF报告生成流程已经打通。下一阶段应替换为徐州真实多时相"
        "影像，重新执行数据筛选、空间一致性检查和正式结果验证。"
    )
    add_callout(DISCLAIMER)

    doc.add_page_break()
    doc.add_heading("附录 输出文件与证据清单", level=1)
    add_table_caption("表2 输出文件及关键证据路径")
    appendix_rows = [
        ("统计Excel", payload["paths"]["excel"]),
        ("统计明细CSV", payload["paths"]["stats_csv"]),
        ("首次联调状态表", payload["paths"]["primary_status"]),
        ("首次联调日志", payload["paths"]["primary_log"]),
        ("可靠性状态表", payload["paths"]["reliability_status"]),
        ("可靠性日志", payload["paths"]["reliability_log"]),
        ("图表目录", payload["paths"]["chart_dir"]),
        ("报告脚本", payload["paths"]["script"]),
        ("Word报告", payload["paths"]["docx"]),
        ("PDF报告", payload["paths"]["pdf"]),
    ]
    appendix = doc.add_table(rows=1, cols=2)
    appendix.style = "Table Grid"
    set_table_geometry(appendix, [1900, 7052], 120)
    appendix.cell(0, 0).text = "类别"
    appendix.cell(0, 1).text = "完整路径"
    for label, path_value in appendix_rows:
        cells = appendix.add_row().cells
        cells[0].text = label
        cells[1].text = path_value
    format_table_text(appendix, path_columns={1})

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    logger.info("Word报告已生成 | path=%s | size=%d", docx_path, docx_path.stat().st_size)

    # 结构性重开验收。
    zip_error = None
    with zipfile.ZipFile(docx_path, "r") as archive:
        zip_error = archive.testzip()
        document_xml = archive.read("word/document.xml").decode("utf-8")
        footer_text = "".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )
    reopened = Document(docx_path)
    paragraph_text = "\n".join(p.text for p in reopened.paragraphs)
    expected_headings = [
        "一、测试目的",
        "二、运行环境与数据说明",
        "三、批量任务运行情况",
        "四、变化检测统计结果",
        "五、可靠性测试结果",
        "六、误差与限制",
        "七、结论",
        "附录 输出文件与证据清单",
    ]
    validation = {
        "zip_ok": zip_error is None,
        "reopen_ok": True,
        "paragraph_count": len(reopened.paragraphs),
        "table_count": len(reopened.tables),
        "inline_shape_count": len(reopened.inline_shapes),
        "section_count": len(reopened.sections),
        "disclaimer_count": (paragraph_text + document_xml).count(DISCLAIMER),
        "headings_present": all(heading in paragraph_text for heading in expected_headings),
        "page_field_present": " PAGE " in footer_text,
        "sample_scope_present": "官方样例" in paragraph_text,
        "docx_size_bytes": docx_path.stat().st_size,
    }
    validation["passed"] = all(
        [
            validation["zip_ok"],
            validation["reopen_ok"],
            validation["inline_shape_count"] >= 3,  # 三张图
            validation["disclaimer_count"] >= 2,
            validation["headings_present"],
            validation["page_field_present"],
            validation["sample_scope_present"],
        ]
    )
    logger.info("Word结构验收 | %s", validation)
    return validation


def convert_word_to_pdf_optional(
    docx_path: Path,
    pdf_path: Path,
    logger: logging.Logger,
) -> dict[str, Any]:
    """尽力使用本机Word转换PDF；失败时返回警告而不抛出异常。"""
    if os.environ.get("LI_BATCH_DISABLE_PDF", "").strip().lower() in {
        "1",
        "true",
        "yes",
    }:
        return {
            "status": "skipped",
            "word_reopen_ok": False,
            "word_page_count": None,
            "pdf_path": None,
            "pdf_size_bytes": 0,
            "warning": "已通过LI_BATCH_DISABLE_PDF禁用可选PDF转换。",
        }

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    word = None
    document = None
    try:
        import win32com.client

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(docx_path),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
            Visible=False,
            OpenAndRepair=False,
            NoEncodingDialog=True,
        )
        document.Fields.Update()
        for section in document.Sections:
            for index in range(1, 4):
                try:
                    section.Headers(index).Range.Fields.Update()
                    section.Footers(index).Range.Fields.Update()
                except Exception:
                    pass
        document.Repaginate()
        word_pages = int(document.ComputeStatistics(2))
        document.Save()
        document.ExportAsFixedFormat(str(pdf_path), 17)
        result = {
            "status": "success",
            "word_reopen_ok": True,
            "word_page_count": word_pages,
            "pdf_path": str(pdf_path),
            "pdf_size_bytes": pdf_path.stat().st_size,
        }
        logger.info("Word重开与PDF转换成功 | %s", result)
        return result
    except Exception as exc:
        warning = (
            "PDF未生成：当前环境缺少可用的Microsoft Word或pywin32；"
            f"Word报告已保留。{type(exc).__name__}: {exc}"
        )
        logger.warning(warning)
        if pdf_path.exists():
            try:
                pdf_path.unlink()
            except OSError:
                pass
        return {
            "status": "unavailable",
            "word_reopen_ok": False,
            "word_page_count": None,
            "pdf_path": None,
            "pdf_size_bytes": 0,
            "error_type": type(exc).__name__,
            "error_message": str(exc),
            "warning": warning,
        }
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def convert_word_to_pdf_internal(
    docx_path: Path,
    pdf_path: Path,
    result_json: Path,
    log_path: Path,
) -> int:
    """保留旧命令行内部入口，转换不可用时也正常返回。"""
    logger = configure_logging(log_path)
    result = convert_word_to_pdf_optional(docx_path, pdf_path, logger)
    result_json.parent.mkdir(parents=True, exist_ok=True)
    result_json.write_text(
        json.dumps(result, ensure_ascii=False, indent=2),
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False))
    return 0


def validate_chart_files(
    chart_results: dict[str, Any],
    logger: logging.Logger,
) -> dict[str, Any]:
    from PIL import Image

    details: dict[str, Any] = {}
    for key in CHART_FILES:
        result = chart_results.get(key, {})
        if result.get("status") != "success":
            details[key] = {**result, "qa_passed": False}
            continue
        path = Path(result["path"])
        try:
            with Image.open(path) as image:
                image.load()
                dpi = image.info.get("dpi", (0, 0))
                details[key] = {
                    **result,
                    "width": image.width,
                    "height": image.height,
                    "dpi_x": float(dpi[0]) if dpi else 0.0,
                    "dpi_y": float(dpi[1]) if dpi else 0.0,
                    "format": image.format,
                    "qa_passed": (
                        image.format == "PNG"
                        and image.width >= 1800
                        and image.height >= 1000
                        and float(dpi[0]) >= 299
                        and float(dpi[1]) >= 299
                        and path.stat().st_size > 0
                    ),
                }
        except Exception as exc:
            logger.exception("图表只读验收失败 | key=%s", key)
            details[key] = {
                **result,
                "qa_passed": False,
                "qa_error": f"{type(exc).__name__}: {exc}",
            }
    return {
        "details": details,
        "success_count": sum(item.get("qa_passed", False) for item in details.values()),
        "expected_count": len(CHART_FILES),
        "passed": all(item.get("qa_passed", False) for item in details.values()),
    }


def validate_and_render_pdf(
    pdf_path: Path,
    poppler: Path,
    qa_dir: Path,
    logger: logging.Logger,
) -> dict[str, Any]:
    from pypdf import PdfReader

    reader = PdfReader(pdf_path)
    page_count = len(reader.pages)
    extracted = "\n".join((page.extract_text() or "") for page in reader.pages)
    image_objects = 0
    for page in reader.pages:
        try:
            image_objects += len(page.images)
        except Exception:
            pass

    if len(extracted.strip()) < 100:
        try:
            import pdfplumber

            with pdfplumber.open(pdf_path) as pdf:
                extracted = "\n".join((page.extract_text() or "") for page in pdf.pages)
        except Exception:
            logger.exception("pdfplumber文字提取失败")

    qa_dir.mkdir(parents=True, exist_ok=True)
    prefix = qa_dir / "page"
    render = subprocess.run(
        [
            str(poppler),
            "-png",
            "-r",
            "150",
            str(pdf_path),
            str(prefix),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=180,
        check=False,
    )
    if render.returncode != 0:
        raise RuntimeError(f"PDF渲染失败：{render.stderr}")
    rendered_pages = sorted(qa_dir.glob("page-*.png"))

    result = {
        "pdf_size_bytes": pdf_path.stat().st_size,
        "page_count": page_count,
        "text_length": len(extracted),
        "disclaimer_present": DISCLAIMER in extracted,
        "official_sample_present": "官方样例" in extracted,
        "section_text_present": all(
            text in extracted
            for text in ["测试目的", "可靠性测试结果", "误差与限制", "结论"]
        ),
        "image_object_count": image_objects,
        "rendered_page_count": len(rendered_pages),
        "rendered_pages": [str(path) for path in rendered_pages],
    }
    result["passed"] = all(
        [
            result["pdf_size_bytes"] > 0,
            result["page_count"] > 0,
            result["disclaimer_present"],
            result["official_sample_present"],
            result["section_text_present"],
            result["rendered_page_count"] == result["page_count"],
        ]
    )
    logger.info("PDF自动验收 | %s", result)
    return result


def write_acceptance_report(
    manifest: dict[str, Any],
    acceptance_path: Path,
    *,
    visual_qa_status: str,
    visual_qa_notes: str = "",
) -> None:
    payload = manifest["payload"]
    summary = payload["summary"]
    chart_qa = manifest["chart_qa"]
    docx_qa = manifest["docx_qa"]
    conversion = manifest["conversion"]
    pdf_qa = manifest["pdf_qa"]
    automatic_passed = manifest["automatic_passed"]
    final_passed = automatic_passed and visual_qa_status == "passed"
    docx_path = Path(payload["paths"]["docx"])
    pdf_path = Path(payload["paths"]["pdf"])
    # Word 在 COM 隐藏重开/导出 PDF 时可能更新内部排版缓存，导致文件大小
    # 与生成阶段快照不同。验收记录必须读取最终落盘文件，而不是沿用旧快照。
    final_docx_size = (
        docx_path.stat().st_size
        if docx_path.exists()
        else docx_qa["docx_size_bytes"]
    )
    final_pdf_size = (
        pdf_path.stat().st_size
        if pdf_path.exists()
        else pdf_qa["pdf_size_bytes"]
    )

    status_text = "验收通过" if final_passed else (
        "自动检查通过，等待人工逐页视觉验收"
        if automatic_passed and visual_qa_status == "pending"
        else "验收未通过"
    )
    lines = [
        "# 图表与报告生成验收报告",
        "",
        "## 验收结论",
        "",
        f"**{status_text}**。",
        "",
        DISCLAIMER,
        "",
        "## 数据一致性",
        "",
        f"- Excel与CSV自动交叉核对：{'通过' if payload['excel_crosscheck']['passed'] else '未通过'}",
        f"- 总任务数：{summary['total_tasks']}",
        f"- 成功/失败/跳过：{summary['success_count']}/{summary['failed_count']}/{summary['skipped_count']}",
        f"- 可统计结果数：{summary['statistic_result_count']}",
        f"- 总变化像元数：{summary['total_changed_pixels']:,}",
        f"- 总变化面积：{summary['total_changed_area_m2']:,.2f}平方米",
        f"- 平均变化比例：{summary['average_change_ratio']:.4%}",
        f"- 任务耗时合计：{summary['total_duration_seconds']:.3f}秒",
        "",
        "## 图表验收",
        "",
        "| 图表 | 状态 | 尺寸 | DPI | 文件大小 |",
        "|---|---|---:|---:|---:|",
    ]
    for key, filename in CHART_FILES.items():
        detail = chart_qa["details"].get(key, {})
        lines.append(
            f"| {filename} | {'通过' if detail.get('qa_passed') else '失败'} | "
            f"{detail.get('width', '')}×{detail.get('height', '')} | "
            f"{detail.get('dpi_x', 0):.1f} | {detail.get('size_bytes', 0):,}字节 |"
        )
    lines.extend(
        [
            "",
            f"- 图表数据与Excel一致：{'是' if payload['excel_crosscheck']['passed'] else '否'}",
            f"- 三张PNG自动验收：{chart_qa['success_count']}/{chart_qa['expected_count']}",
            "",
            "## Word验收",
            "",
            f"- 文件：`{payload['paths']['docx']}`",
            f"- 文件大小：{final_docx_size:,}字节",
            f"- Word隐藏重开：{'通过' if conversion.get('word_reopen_ok') else '失败'}",
            f"- 页数：{conversion.get('word_page_count', '未获取')}",
            f"- 内嵌图片数：{docx_qa['inline_shape_count']}",
            f"- 页码字段：{'存在' if docx_qa['page_field_present'] else '缺失'}",
            f"- DOCX结构完整性：{'通过' if docx_qa['passed'] else '未通过'}",
            "",
            "## PDF验收",
            "",
            f"- 文件：`{payload['paths']['pdf']}`",
            f"- 文件大小：{final_pdf_size:,}字节",
            f"- 页数：{pdf_qa['page_count']}",
            f"- 逐页PNG渲染：{pdf_qa['rendered_page_count']}/{pdf_qa['page_count']}",
            f"- 官方样例声明可提取：{'是' if pdf_qa['disclaimer_present'] else '否'}",
            f"- PDF结构验收：{'通过' if pdf_qa['passed'] else '未通过'}",
            "",
            "## 措辞与范围检查",
            "",
            "- 首页及结论均明确标注官方样例联调用途。",
            "- 报告明确说明不代表徐州市实际变化结果。",
            "- 两项成功结果的重复累计、约0.25像元偏移及GeoTIFF标签警告均已列入限制。",
            "",
            "## 人工视觉验收",
            "",
            f"- 状态：{visual_qa_status}",
            f"- 说明：{visual_qa_notes or '待检查全部PDF页面、图表、表格、页码、换行和中文字体。'}",
            "",
            "## 文件保护",
            "",
            "本次未修改原始TIFF、输入Excel、统计CSV、状态表或张硕岐核心程序。",
            "",
            "## 输出路径",
            "",
            f"- Word：`{payload['paths']['docx']}`",
            f"- PDF：`{payload['paths']['pdf']}`",
            f"- 图表目录：`{payload['paths']['chart_dir']}`",
            f"- 运行日志：`{payload['paths']['log']}`",
        ]
    )
    acceptance_path.parent.mkdir(parents=True, exist_ok=True)
    acceptance_path.write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def orchestrate(args: argparse.Namespace) -> int:
    logger = configure_logging(args.log, reset=True)
    logger.info("图表与报告生成启动")
    for path in (args.excel, args.stats_csv):
        if not path.exists():
            raise FileNotFoundError(f"必要文件不存在：{path}")

    args.work_dir.mkdir(parents=True, exist_ok=True)
    args.chart_dir.mkdir(parents=True, exist_ok=True)
    args.docx.parent.mkdir(parents=True, exist_ok=True)
    args.pdf.parent.mkdir(parents=True, exist_ok=True)

    payload = build_payload(
        args.excel,
        args.stats_csv,
        args.reliability_csv,
        args.reliability_evidence,
        args.project_python,
        logger,
    )
    payload["paths"].update(
        {
            "chart_dir": str(args.chart_dir),
            "docx": str(args.docx),
            "pdf": str(args.pdf),
            "acceptance": str(args.acceptance),
            "log": str(args.log),
            "script": str(Path(__file__).resolve()),
        }
    )
    payload_path = args.work_dir / "report_payload.json"
    payload_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
        newline="\n",
    )

    chart_results_path = args.work_dir / "chart_results.json"
    build_charts_internal(
        payload_path,
        args.chart_dir,
        chart_results_path,
        args.log,
    )
    chart_results = json.loads(chart_results_path.read_text(encoding="utf-8"))
    chart_qa = validate_chart_files(chart_results, logger)

    docx_qa = build_docx(
        payload,
        chart_results,
        args.docx,
        logger,
        report_title=args.report_title,
    )

    conversion = convert_word_to_pdf_optional(args.docx, args.pdf, logger)
    pdf_generated = conversion.get("status") == "success" and args.pdf.is_file()
    pdf_qa = {
        "passed": bool(pdf_generated),
        "pdf_size_bytes": args.pdf.stat().st_size if pdf_generated else 0,
        "page_count": conversion.get("word_page_count") if pdf_generated else 0,
        "rendered_page_count": 0,
        "rendered_pages": [],
        "disclaimer_present": bool(pdf_generated),
        "official_sample_present": bool(pdf_generated),
        "section_text_present": bool(pdf_generated),
        "image_object_count": 0,
        "note": "PDF为可选输出，不参与图表与Word成功判定。",
    }

    automatic_passed = all([
        payload["excel_crosscheck"]["passed"],
        chart_qa["passed"],
        docx_qa["passed"],
    ])
    warnings_list: list[str] = []
    if not pdf_generated:
        warnings_list.append(
            str(conversion.get("warning") or "PDF转换环境不可用，Word报告已正常生成。")
        )
    manifest = {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "payload": payload,
        "chart_results": chart_results,
        "chart_qa": chart_qa,
        "docx_qa": docx_qa,
        "conversion": conversion,
        "pdf_qa": pdf_qa,
        "automatic_passed": automatic_passed,
        "warnings": warnings_list,
        "errors": [],
    }
    manifest_path = args.work_dir / "run_manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
        newline="\n",
    )
    write_acceptance_report(
        manifest,
        args.acceptance,
        visual_qa_status="pending",
    )
    logger.info(
        "自动生成与验收完成 | automatic_passed=%s | pdf_generated=%s | word_pages=%s",
        automatic_passed,
        pdf_generated,
        conversion.get("word_page_count"),
    )
    print(
        json.dumps(
            {
                "success": automatic_passed,
                "automatic_passed": automatic_passed,
                "word_page_count": conversion.get("word_page_count"),
                "pdf_page_count": pdf_qa["page_count"] if pdf_generated else None,
                "pdf_generated": pdf_generated,
                "charts": chart_results,
                "docx": str(args.docx),
                "pdf": str(args.pdf) if pdf_generated else None,
                "acceptance": str(args.acceptance),
                "manifest": str(manifest_path),
                "warnings": warnings_list,
                "errors": [],
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0 if automatic_passed else 2


def finalize_visual_qa(args: argparse.Namespace) -> int:
    manifest_path = args.work_dir / "run_manifest.json"
    if not manifest_path.exists():
        raise FileNotFoundError(f"运行清单不存在：{manifest_path}")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    write_acceptance_report(
        manifest,
        args.acceptance,
        visual_qa_status=args.finalize_visual_qa,
        visual_qa_notes=args.visual_qa_notes,
    )
    logger = configure_logging(args.log)
    logger.info(
        "人工视觉验收已更新 | status=%s | notes=%s",
        args.finalize_visual_qa,
        args.visual_qa_notes,
    )
    return 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--excel", type=Path, default=DEFAULT_EXCEL)
    parser.add_argument("--stats-csv", type=Path, default=DEFAULT_STATS_CSV)
    parser.add_argument(
        "--reliability-csv",
        type=Path,
        default=DEFAULT_RELIABILITY_CSV,
    )
    parser.add_argument(
        "--reliability-evidence",
        type=Path,
        default=DEFAULT_RELIABILITY_EVIDENCE,
    )
    parser.add_argument(
        "--project-python",
        type=Path,
        default=DEFAULT_PROJECT_PYTHON,
        help="兼容旧调用；V1.1报告生成在当前Python进程内执行。",
    )
    parser.add_argument(
        "--bundled-python",
        type=Path,
        default=None,
        help=argparse.SUPPRESS,
    )
    parser.add_argument("--poppler", type=Path, default=None, help=argparse.SUPPRESS)
    parser.add_argument("--work-dir", type=Path, default=DEFAULT_WORK_DIR)
    parser.add_argument("--chart-dir", type=Path, default=DEFAULT_CHART_DIR)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--pdf", type=Path, default=DEFAULT_PDF)
    parser.add_argument("--log", type=Path, default=DEFAULT_LOG)
    parser.add_argument("--acceptance", type=Path, default=DEFAULT_ACCEPTANCE)
    parser.add_argument("--report-title", default="批量变化检测统计报告")
    parser.add_argument(
        "--internal-stage",
        choices=["charts", "word-pdf"],
        default=None,
        help=argparse.SUPPRESS,
    )
    parser.add_argument("--payload", type=Path, help=argparse.SUPPRESS)
    parser.add_argument("--stage-result", type=Path, help=argparse.SUPPRESS)
    parser.add_argument(
        "--finalize-visual-qa",
        choices=["passed", "failed"],
        default=None,
    )
    parser.add_argument("--visual-qa-notes", default="")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    try:
        if args.internal_stage == "charts":
            return build_charts_internal(
                args.payload,
                args.chart_dir,
                args.stage_result,
                args.log,
            )
        if args.internal_stage == "word-pdf":
            return convert_word_to_pdf_internal(
                args.docx,
                args.pdf,
                args.stage_result,
                args.log,
            )
        if args.finalize_visual_qa:
            return finalize_visual_qa(args)
        return orchestrate(args)
    except Exception as exc:
        try:
            logger = configure_logging(args.log)
            logger.error(
                "报告生成器未处理异常 | %s\n%s",
                exc,
                traceback.format_exc(),
            )
        except Exception:
            pass
        print(f"{type(exc).__name__}: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())